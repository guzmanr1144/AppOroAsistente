import os, json, ast, re, warnings, copy
warnings.filterwarnings("ignore", category=DeprecationWarning)
import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
import PyPDF2
from fpdf import FPDF
try:
    import fitz
    PYMUPDF_OK = True
except ImportError:
    PYMUPDF_OK = False
from io import BytesIO
from datetime import datetime
import pytz

st.set_page_config(page_title="Oro Asistente", page_icon="🏆", layout="centered", initial_sidebar_state="collapsed")

# ══════════════════════════════════════════════════════════════
# CSS — cacheado por tema
# ══════════════════════════════════════════════════════════════
@st.cache_data(show_spinner=False)
def _get_all_css(tema_key="noche"):
    _T = {
        "noche": {
            "bg1":"#0a0e1a","bg2":"#0d1222","bg3":"#10162a",
            "card":"#141828","card2":"#1a2035",
            "borde":"#2a3560","borde2":"#364480",
            "acento1":"#6b83f8","acento2":"#8fa0ff","acento3":"#4f6ef7",
            "titulo_grad":"linear-gradient(135deg,#6b83f8,#a78bfa,#6b83f8)",
            "texto":"#e8ecff","texto2":"#a8b8f0","texto3":"#5060a0",
            "sombra":"rgba(107,131,248,0.2)","sombra2":"rgba(107,131,248,0.08)",
        },
        "carbon": {
            "bg1":"#111418","bg2":"#161b22","bg3":"#1c2330",
            "card":"#1e2530","card2":"#242e3d",
            "borde":"#2d3f55","borde2":"#3a5070",
            "acento1":"#10b981","acento2":"#34d399","acento3":"#6ee7b7",
            "titulo_grad":"linear-gradient(135deg,#10b981,#06b6d4,#10b981)",
            "texto":"#d8f0e8","texto2":"#90c8b0","texto3":"#3a6050",
            "sombra":"rgba(16,185,129,0.2)","sombra2":"rgba(16,185,129,0.08)",
        },
        "cosmos": {
            "bg1":"#0d0818","bg2":"#120d22","bg3":"#18102c",
            "card":"#1a1030","card2":"#201540",
            "borde":"#3a2060","borde2":"#502888",
            "acento1":"#a78bfa","acento2":"#c4b0ff","acento3":"#7c3aed",
            "titulo_grad":"linear-gradient(135deg,#a78bfa,#f472b6,#a78bfa)",
            "texto":"#f0e8ff","texto2":"#c0a8f0","texto3":"#705898",
            "sombra":"rgba(167,139,250,0.2)","sombra2":"rgba(167,139,250,0.08)",
        },
    }
    t = _T.get(tema_key, _T["noche"])
    return f"""<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&family=JetBrains+Mono:wght@400;500;600&display=swap');
*{{box-sizing:border-box}}
html,body,[class*="css"]{{font-family:'Inter',sans-serif!important;-webkit-tap-highlight-color:transparent}}
.stApp{{background:linear-gradient(145deg,{t['bg2']} 0%,{t['bg1']} 40%,{t['bg3']} 100%)!important;min-height:100vh}}
.main .block-container{{padding:.8rem .9rem 5rem .9rem!important;max-width:460px!important;margin:0 auto!important;background:transparent!important}}
#MainMenu,footer,header{{visibility:hidden}}[data-testid="stToolbar"]{{display:none}}
.oro-header{{text-align:center;padding:1.6rem 0 .5rem}}
.oro-logo-wrap{{position:relative;display:inline-block;margin-bottom:.3rem}}
.oro-logo{{font-size:2.8rem;display:block;filter:drop-shadow(0 4px 12px {t['sombra']});animation:float 4s ease-in-out infinite}}
.oro-logo-ring{{position:absolute;inset:-8px;border:2px solid {t['acento1']};border-radius:50%;opacity:.2;animation:spin 8s linear infinite}}
@keyframes float{{0%,100%{{transform:translateY(0)}}50%{{transform:translateY(-5px)}}}}
@keyframes spin{{to{{transform:rotate(360deg)}}}}
@keyframes fadeUp{{from{{opacity:0;transform:translateY(8px)}}to{{opacity:1;transform:translateY(0)}}}}
@keyframes shimmer{{0%,100%{{background-position:0% 50%}}50%{{background-position:100% 50%}}}}
.oro-title{{font-size:1.9rem;font-weight:900;background:{t['titulo_grad']}!important;background-size:200%!important;-webkit-background-clip:text!important;-webkit-text-fill-color:transparent!important;background-clip:text!important;letter-spacing:-.03em;animation:shimmer 5s ease infinite}}
.oro-badge{{display:inline-flex;align-items:center;gap:.3rem;background:{t['card']};border:1px solid {t['borde']};border-radius:20px;padding:.2rem .8rem;font-size:.65rem;color:{t['acento1']};font-weight:700;letter-spacing:.06em;margin-top:.3rem;box-shadow:0 2px 8px {t['sombra2']}}}
.file-badge{{display:flex;align-items:center;gap:.75rem;background:{t['card']};border:1px solid {t['borde']};border-radius:18px;padding:.85rem 1rem;margin:.5rem 0;animation:fadeUp .3s ease;box-shadow:0 4px 16px {t['sombra2']}}}
.file-icon{{font-size:1.6rem;flex-shrink:0}}
.file-info-name{{color:{t['texto']}!important;font-weight:700;font-size:.85rem;word-break:break-all;line-height:1.3}}
.file-info-stats{{color:{t['texto3']}!important;font-size:.7rem;margin-top:.15rem;display:flex;gap:.4rem;flex-wrap:wrap}}
.stat-chip{{background:{t['bg2']};border:1px solid {t['borde']};border-radius:8px;padding:.05rem .4rem;font-size:.65rem;color:{t['acento1']};font-weight:600}}
.stButton>button{{background:{t['card']}!important;color:{t['texto2']}!important;border:1.5px solid {t['borde']}!important;border-radius:12px!important;font-weight:600!important;font-size:.84rem!important;min-height:3rem!important;width:100%!important;transition:all .15s ease!important;font-family:'Inter',sans-serif!important;box-shadow:0 2px 6px {t['sombra2']}!important}}
.stButton>button:hover{{border-color:{t['acento1']}!important;color:{t['acento2']}!important;background:{t['bg2']}!important;box-shadow:0 4px 14px {t['sombra']}!important;transform:translateY(-1px)!important}}
.stButton>button:active{{transform:scale(.97)!important;box-shadow:none!important}}
.btn-analizar>button{{background:linear-gradient(135deg,{t['acento1']},{t['acento3']})!important;color:white!important;border:none!important;font-weight:700!important;box-shadow:0 4px 14px {t['sombra']}!important}}
.btn-analizar>button:hover{{filter:brightness(1.08)!important;box-shadow:0 6px 20px {t['sombra']}!important}}
.btn-evaluar>button{{background:linear-gradient(135deg,#059669,#0891b2)!important;color:white!important;border:none!important;font-weight:700!important;box-shadow:0 4px 14px rgba(5,150,105,.2)!important}}
.btn-evaluar>button:hover{{filter:brightness(1.08)!important}}
.btn-peligro>button{{background:linear-gradient(135deg,#dc2626,#e11d48)!important;color:white!important;border:none!important;font-weight:600!important}}
[data-testid=\"stDownloadButton\"]>button{{background:linear-gradient(135deg,{t['acento1']},{t['acento3']})!important;color:white!important;border:none!important;border-radius:12px!important;font-weight:700!important;height:2.8rem!important;width:100%!important;box-shadow:0 3px 10px {t['sombra']}!important;transition:all .15s!important}}
[data-testid=\"stDownloadButton\"]>button:hover{{filter:brightness(1.08)!important;box-shadow:0 5px 16px {t['sombra']}!important;transform:translateY(-1px)!important}}
.nav-tab-activo>button{{background:linear-gradient(135deg,{t['acento1']},{t['acento3']})!important;color:white!important;border:none!important;font-weight:700!important;box-shadow:0 3px 10px {t['sombra']}!important}}
.nav-tab-inactivo>button{{background:{t['card']}!important;color:{t['texto3']}!important;border:1.5px solid {t['borde']}!important;font-weight:500!important}}
[data-testid=\"stFileUploader\"]{{background:transparent!important;border:none!important}}
[data-testid=\"stFileUploader\"]>div{{background:{t['card2']}!important;border:2px dashed {t['borde2']}!important;border-radius:20px!important;padding:1.3rem!important;box-shadow:0 4px 16px {t['sombra']}!important;transition:border-color .3s!important}}
[data-testid=\"stFileUploader\"]>div:hover{{border-color:{t['acento1']}!important}}
[data-testid=\"stFileUploader\"] label{{color:{t['acento1']}!important;font-weight:600!important;font-size:.92rem!important}}
.summary-card{{background:{t['card']};border:1px solid {t['borde']};border-left:4px solid {t['acento1']};border-radius:18px;padding:1.1rem 1.2rem;margin:.7rem 0;color:{t['texto2']}!important;line-height:1.75;font-size:.88rem;box-shadow:0 4px 20px {t['sombra2']};animation:fadeUp .4s ease}}
.summary-card-title{{color:{t['acento2']}!important;font-size:1rem;font-weight:800;margin-bottom:.5rem;display:flex;align-items:center;gap:.4rem}}
.metrics-grid{{display:grid;grid-template-columns:1fr 1fr;gap:.5rem;margin:.6rem 0}}
.metric-pill{{background:{t['card']};border:1px solid {t['borde']};border-radius:14px;padding:.75rem 1rem;text-align:center;box-shadow:0 2px 8px {t['sombra2']};transition:transform .2s,box-shadow .2s}}
.metric-pill:hover{{transform:translateY(-2px);box-shadow:0 6px 16px {t['sombra']}}}
.metric-pill-label{{color:{t['texto3']}!important;font-size:.62rem;font-weight:700;text-transform:uppercase;letter-spacing:.08em}}
.metric-pill-value{{color:{t['acento2']}!important;font-size:1.1rem;font-weight:800;margin-top:.2rem;font-family:'JetBrains Mono',monospace}}
.tags-wrap{{display:flex;flex-wrap:wrap;gap:.3rem;margin:.5rem 0}}
.tag{{background:{t['bg2']};color:{t['acento2']}!important;border:1px solid {t['borde']};border-radius:20px;padding:.28rem .7rem;font-size:.7rem;font-weight:600;transition:all .2s}}
.tag:hover{{background:{t['borde']};border-color:{t['acento1']}}}
.hallazgo-card{{background:linear-gradient(135deg,{t['bg2']},{t['card']});border:1px solid {t['borde']};border-left:4px solid {t['acento1']};border-radius:14px;padding:.85rem 1rem;color:{t['texto2']}!important;font-size:.82rem;margin:.6rem 0;line-height:1.65;box-shadow:0 2px 8px {t['sombra2']}}}
.info-box{{background:#f0fdf4;border:1px solid #86efac;border-radius:12px;padding:.75rem 1rem;color:#166534;font-size:.83rem;margin:.5rem 0;display:flex;align-items:center;gap:.5rem}}
.warn-box{{background:#fffbeb;border:1px solid #fcd34d;border-radius:12px;padding:.75rem 1rem;color:#92400e;font-size:.83rem;margin:.5rem 0;display:flex;align-items:center;gap:.5rem}}
.cambio-item{{background:{t['bg2']};border:1px solid {t['borde']};border-radius:10px;padding:.5rem .8rem;margin:.2rem 0;font-size:.78rem;font-family:'JetBrains Mono',monospace;display:flex;align-items:center;gap:.45rem;color:{t['texto']}}}
.cambio-num{{color:{t['texto3']};font-size:.66rem;min-width:1rem}}.cambio-arrow{{color:#d97706}}
[data-testid=\"stExpander\"]{{background:{t['card']}!important;border:1px solid {t['borde']}!important;border-radius:14px!important;box-shadow:0 2px 8px {t['sombra2']}!important;overflow:hidden!important}}
[data-testid=\"stChatInput\"] textarea{{background:{t['card2']}!important;border:2px solid {t['borde2']}!important;border-radius:16px!important;color:{t['texto']}!important;font-family:'Inter',sans-serif!important;font-size:.92rem!important;box-shadow:0 2px 8px {t['sombra2']}!important}}
[data-testid=\"stChatInput\"] textarea:focus{{border-color:{t['acento1']}!important;box-shadow:0 0 0 3px {t['sombra']}!important}}[data-testid=\"stBottom\"]{{background:linear-gradient(transparent,{t['bg1']} 40%)!important;padding-top:1rem!important}}
[data-testid=\"stChatMessageContent\"]{{font-size:.88rem!important;line-height:1.65!important}}
[data-testid=\"stChatMessage\"]{{background:{t['card']}!important;border:1px solid {t['borde']}!important;border-radius:14px!important;padding:.7rem!important;margin:.3rem 0!important;box-shadow:0 2px 6px {t['sombra2']}!important}}
.stTextInput>div>div>input{{background:{t['card']}!important;border:1.5px solid {t['borde']}!important;border-radius:12px!important;color:{t['texto']}!important;font-family:'Inter',sans-serif!important;font-size:.88rem!important;padding:.6rem .9rem!important}}
.stTextInput>div>div>input:focus{{border-color:{t['acento1']}!important;box-shadow:0 0 0 3px {t['sombra']}!important}}
.stSelectbox>div>div{{background:{t['card']}!important;border:1.5px solid {t['borde']}!important;border-radius:12px!important;color:{t['texto']}!important}}
.oro-divider{{height:1px;background:linear-gradient(90deg,transparent,{t['borde2']},transparent);margin:.9rem 0}}
.section-title{{color:{t['texto']}!important;font-size:.95rem;font-weight:700;margin:.9rem 0 .4rem;display:flex;align-items:center;gap:.4rem}}
.chat-placeholder{{text-align:center;padding:1.2rem .8rem;background:{t['card']};border:1.5px dashed {t['borde2']};border-radius:18px;margin:.4rem 0;box-shadow:0 2px 8px {t['sombra2']}}}
.chip{{display:inline-block;background:{t['bg2']};border:1px solid {t['borde']};border-radius:20px;padding:.22rem .65rem;font-size:.68rem;color:{t['acento1']};font-weight:600;margin:.15rem;cursor:default;transition:all .2s}}
.chip:hover{{background:{t['borde']};border-color:{t['acento1']}}}[data-testid=\"stChatInputContainer\"]{{padding:.6rem 0!important}}[data-testid=\"stChatInput\"]{{margin-top:.4rem!important}}
.empty-state{{text-align:center;padding:3rem 1rem;animation:fadeUp .5s ease}}
.empty-icon{{font-size:3.5rem;margin-bottom:.8rem;filter:drop-shadow(0 4px 12px {t['sombra']})}}
.empty-title{{color:{t['texto']};font-size:1rem;font-weight:700}}
.empty-hint{{color:{t['texto3']};font-size:.78rem;margin-top:.4rem;line-height:1.7}}
.format-badges{{display:flex;justify-content:center;gap:.4rem;margin-top:.8rem;flex-wrap:wrap}}
.format-badge{{background:{t['bg2']};border:1px solid {t['borde']};border-radius:8px;padding:.22rem .6rem;color:{t['texto3']};font-size:.7rem;font-family:'JetBrains Mono',monospace;font-weight:600}}
.oro-footer{{text-align:center;font-size:.68rem;color:{t['texto3']};padding:.5rem 0;opacity:.7}}
[data-testid=\"stSidebar\"]{{background:linear-gradient(180deg,{t['card']},{t['bg2']})!important;border-right:1px solid {t['borde']}!important}}
[data-testid=\"stSidebar\"] .stMarkdown,
[data-testid=\"stSidebar\"] .stCaption{{color:{t['texto2']}!important}}
</style>"""


# ══════════════════════════════════════════════════════════════
# SESSION STATE
# ══════════════════════════════════════════════════════════════
_defaults = {
    "texto_extraido":"","nombre_archivo":"","archivo_bytes":None,"resumen_data":None,
    "historial_chat":[],"cambios_aplicados":None,"archivo_tipo":"","lista_cambios":[],
    "texto_modificado":"","generando_resumen":False,"resumen_error":False,
    "preview_cambio":None,"edicion_counter":0,"texto_corregido":"",
    "guia_paso":0,"guia_vista":False,"ejecutar_evaluacion":False,
    "tema":"noche","modo_entrada":"archivo",
    "imagen_archivo_bytes":None,"imagen_archivo_nombre":"","imagen_archivo_mime":"",
    "historial_versiones":[],"buscar_query":"",
    "resultado_evaluacion":None,
    "scroll_to":None,
    "idioma":"es",
    "vista_activa":None,
}
for k,v in _defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

st.markdown(_get_all_css(st.session_state.get("tema","noche")), unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════
# TRADUCCIONES
# ══════════════════════════════════════════════════════════════
_TXT = {
    "es": {
        "analizar":"⚡ Analizar","evaluar":"🔍 Evaluar","ver_doc":"👁 Ver documento",
        "analizar_doc":"⚡ Analizar documento","evaluar_doc":"🔍 Evaluar documento ahora",
        "exportar":"📥 Exportar informe","regen":"🔄 Regenerar resumen",
        "confirmar":"✅ Confirmar","cancelar":"❌ Cancelar",
        "descargar_corregido":"📥 Descargar corregido",
        "limpiar_cambios":"🗑️ Limpiar todos los cambios",
        "chat_placeholder":"✍️ Escribe un cambio o una pregunta...",
        "chat_titulo":"Conversa sobre el documento",
        "chat_hint":"Edita, pregunta o pide cambios en lenguaje natural",
        "chip1":"✏️ cambia X por Y","chip2":"➕ agrega dato a persona",
        "chip3":"❓ ¿cuántos hay?","chip4":"📝 resume en 3 puntos",
        "analizando":"🧠 Analizando...","evaluando":"🔎 Evaluando calidad...",
        "procesando":"🔍 Procesando...","pensando":"🤔 Pensando...",
        "cargando":"📖 Cargando...","interpretando":"🧠 Leyendo...",
        "reintentar":"🔄 Reintentar","cerrar_eval":"✕ Cerrar evaluación",
        "hallazgo":"💡 Hallazgo:","recomendacion":"💡 Recomendación:",
        "antes":"Antes","despues":"Después",
        "version":"↩️ Deshacer último cambio",
        "idioma_doc":"El programa responde en español sin importar el idioma del documento.",
        "subir":"📎 Sube tu archivo","foto":"📷 Foto de documento",
        "interpretar":"🔍 Interpretar","archivo":"📎 Archivo",
        "palabras":"palabras","cambios":"cambio(s)","versiones":"versión(es)",
        "no_encontrado":"no encontrado en el documento",
        "listo_analizar":"Toca ⚡ Analizar para generar el resumen inteligente",
        "sin_problemas":"¡Sin problemas detectados! El documento está bien 🎉",
        "prompt_idioma":"Responde SIEMPRE en español, sin importar el idioma del documento.",
        "word":"📄 Word","excel":"📊 Excel","pdf":"📕 PDF",
        "word_corregido":"📄 Word corregido","excel_corregido":"📊 Excel corregido",
        "pdf_corregido":"📕 PDF corregido","exportar_word":"📄 Exportar como Word",
        "cambiar_tema":"🎨 Tema","cambiar_idioma":"🌐 EN",
    },
    "en": {
        "analizar":"⚡ Analyze","evaluar":"🔍 Evaluate","ver_doc":"👁 View document",
        "analizar_doc":"⚡ Analyze document","evaluar_doc":"🔍 Evaluate document now",
        "exportar":"📥 Export report","regen":"🔄 Regenerate summary",
        "confirmar":"✅ Confirm","cancelar":"❌ Cancel",
        "descargar_corregido":"📥 Download corrected",
        "limpiar_cambios":"🗑️ Clear all changes",
        "chat_placeholder":"✍️ Write a change or ask a question...",
        "chat_titulo":"Chat about the document",
        "chat_hint":"Edit, ask questions or request changes in natural language",
        "chip1":"✏️ change X to Y","chip2":"➕ add data to person",
        "chip3":"❓ how many are there?","chip4":"📝 summarize in 3 points",
        "analizando":"🧠 Analyzing...","evaluando":"🔎 Evaluating quality...",
        "procesando":"🔍 Processing...","pensando":"🤔 Thinking...",
        "cargando":"📖 Loading...","interpretando":"🧠 Reading...",
        "reintentar":"🔄 Retry","cerrar_eval":"✕ Close evaluation",
        "hallazgo":"💡 Finding:","recomendacion":"💡 Recommendation:",
        "antes":"Before","despues":"After",
        "version":"↩️ Undo last change",
        "idioma_doc":"The program responds in English regardless of the document language.",
        "subir":"📎 Upload your file","foto":"📷 Photo of document",
        "interpretar":"🔍 Interpret","archivo":"📎 File",
        "palabras":"words","cambios":"change(s)","versiones":"version(s)",
        "no_encontrado":"not found in document",
        "listo_analizar":"Tap ⚡ Analyze to generate the smart summary",
        "sin_problemas":"No issues detected! The document looks good 🎉",
        "prompt_idioma":"Always respond in English, regardless of the document language.",
        "word":"📄 Word","excel":"📊 Excel","pdf":"📕 PDF",
        "word_corregido":"📄 Corrected Word","excel_corregido":"📊 Corrected Excel",
        "pdf_corregido":"📕 Corrected PDF","exportar_word":"📄 Export as Word",
        "cambiar_tema":"🎨 Theme","cambiar_idioma":"🌐 ES",
    },
}

def T(key):
    """Get translated text."""
    lang = st.session_state.get("idioma","es")
    return _TXT.get(lang,_TXT["es"]).get(key, _TXT["es"].get(key,""))

# ══════════════════════════════════════════════════════════════
# GEMINI
# ══════════════════════════════════════════════════════════════
try:
    LLAVE_GEMINI = st.secrets["LLAVE_GEMINI"]
    genai.configure(api_key=LLAVE_GEMINI)
except Exception as e:
    st.error(f"🔑 Error configurando la IA: {e}"); st.stop()

MODELOS_FALLBACK = ["gemini-3.1-flash-lite-preview","gemini-3.1-flash-preview","gemini-3.1-pro-preview"]

def llamar_ia(prompt, es_json=False):
    for modelo in MODELOS_FALLBACK:
        try:
            model = genai.GenerativeModel(modelo)
            resp = model.generate_content(prompt)
            texto = resp.text
            if es_json:
                return extraer_json_seguro(texto, es_lista=texto.strip().startswith("["))
            return texto
        except Exception:
            continue
    return None

# ══════════════════════════════════════════════════════════════
# SIDEBAR
# ══════════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown(f"<div style='text-align:center;padding:.5rem 0'><span style='font-size:1.5rem'>🏆</span><div style='font-size:.9rem;font-weight:800;color:#34d399'>Oro Asistente</div></div>", unsafe_allow_html=True)
    st.markdown("---")
    st.markdown("**🎨 Tema**")
    tema_opc = {"☀️ Claro":"claro","🌸 Aurora":"aurora","🌿 Menta":"menta","🌅 Sol":"sol","🌹 Rose":"rose","🌑 Noche":"noche","⬛ Carbón":"carbon","🌌 Cosmos":"cosmos"}
    sel = st.selectbox("Tema", list(tema_opc.keys()),
        index=list(tema_opc.values()).index(st.session_state.get("tema","noche")),
        label_visibility="collapsed")
    if tema_opc[sel] != st.session_state.tema:
        st.session_state.tema = tema_opc[sel]; st.rerun()
    st.markdown("---")
    paso_sb = st.session_state.get("guia_paso",0)
    if paso_sb > 0 and not st.session_state.get("guia_vista",False):
        guias_sb = {
            1:("🎉","Paso 1 — Analiza","Tu archivo está listo.\n\nToca **⚡ Analizar** para que la IA extraiga el resumen, métricas y puntos clave."),
            2:("📊","Paso 2 — Revisa","El resumen está arriba.\n\nDescárgalo como Word, Excel o PDF con los botones de exportación."),
            3:("💬","Paso 3 — Edita","Usa el chat para editar:\n• *cambia X por Y*\n• *agrega el teléfono a Juan*\n• *¿cuántas personas hay?*"),
        }
        if paso_sb in guias_sb:
            ico_sb,tit_sb,desc_sb = guias_sb[paso_sb]
            st.markdown(f"**{ico_sb} {tit_sb}**")
            st.info(desc_sb)
            c1sb,c2sb = st.columns(2)
            with c1sb:
                if st.button("👍 Ok",use_container_width=True,key="guia_ok"):
                    st.session_state.guia_paso = 0 if paso_sb>=3 else paso_sb+1
                    if paso_sb>=3: st.session_state.guia_vista=True
                    st.rerun()
            with c2sb:
                if st.button("✕ Saltar",use_container_width=True,key="guia_skip"):
                    st.session_state.guia_vista=True; st.session_state.guia_paso=0; st.rerun()
            st.caption(f"Paso {paso_sb} de 3")
    elif st.session_state.get("texto_extraido"):
        st.markdown("**💡 Comandos útiles**")
        st.markdown(f"""<div style='font-size:.78rem;color:#6ee7b7;line-height:2'>
            ✏️ cambia X por Y<br>
            ➕ agrega dato a persona<br>
            🔍 ¿dónde aparece X?<br>
            📝 resume en 3 puntos<br>
            📊 ¿cuántos registros hay?
        </div>""", unsafe_allow_html=True)
        if st.session_state.get("historial_versiones"):
            st.markdown("---")
            st.markdown("**⏮ Versiones**")
            st.caption(f"{len(st.session_state.historial_versiones)} versión(es) guardada(s)")
            if st.button(T("version"), use_container_width=True):
                v = st.session_state.historial_versiones.pop()
                st.session_state.texto_corregido = v["texto"]
                st.session_state.cambios_aplicados = v["bytes"]
                if st.session_state.lista_cambios:
                    st.session_state.lista_cambios.pop()
                st.session_state.resumen_data = None
                st.rerun()
    else:
        st.markdown("**👋 Bienvenido**")
        st.info("Sube un archivo Word, Excel o PDF — o una **foto** de un documento — para empezar.")
    st.markdown("---")
    st.caption("Oro Asistente v3.1")

# ══════════════════════════════════════════════════════════════
# TOP BAR — tema izquierda, idioma derecha
# ══════════════════════════════════════════════════════════════
_col_tema, _col_mid, _col_lang = st.columns([2, 3, 1])

with _col_tema:
    _temas_map = {"🌑 Noche":"noche","⬛ Carbón":"carbon","🌌 Cosmos":"cosmos"}
    _tema_actual = st.session_state.get("tema","noche")
    _tema_label = next((k for k,v in _temas_map.items() if v==_tema_actual), "☀️ Claro")
    _nuevo_tema = st.selectbox(
        "tema", list(_temas_map.keys()),
        index=list(_temas_map.keys()).index(_tema_label),
        label_visibility="collapsed", key="tema_sel")
    if _temas_map[_nuevo_tema] != _tema_actual:
        st.session_state.tema = _temas_map[_nuevo_tema]; st.rerun()

with _col_lang:
    _lang_actual = st.session_state.get("idioma","es")
    _lang_label = "🌐 EN" if _lang_actual=="es" else "🌐 ES"
    if st.button(_lang_label, key="btn_lang", use_container_width=True):
        st.session_state.idioma = "en" if _lang_actual=="es" else "es"
        st.rerun()

# ── Header ──
st.markdown("""
<div class="oro-header">
    <div class="oro-logo-wrap">
        <div class="oro-logo-ring"></div>
        <span class="oro-logo">🏆</span>
    </div>
    <div class="oro-title">Oro Asistente</div>
    <div class="oro-badge">✦ ANALIZA &nbsp;·&nbsp; EDITA &nbsp;·&nbsp; EXPORTA ✦</div>
</div>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════
# FUNCIONES UTILITARIAS
# ══════════════════════════════════════════════════════════════
def extraer_json_seguro(texto, es_lista=False):
    t = texto.replace("```json","").replace("```","").strip()
    c1,c2 = ("[","]") if es_lista else ("{","}")
    ini=t.find(c1); fin=t.rfind(c2)+1
    if ini!=-1 and fin>0:
        try: return json.loads(t[ini:fin],strict=False)
        except:
            try: return ast.literal_eval(t[ini:fin])
            except: pass
    return None

def _scroll_to(anchor_id):
    """Hace scroll suave al elemento con ese id."""
    st.markdown(f'''<script>
        window.parent.document.getElementById("{anchor_id}") &&
        window.parent.document.getElementById("{anchor_id}").scrollIntoView({{behavior:"smooth",block:"start"}});
    </script>''', unsafe_allow_html=True)

def guardar_version(texto, bytes_doc):
    """Guarda snapshot antes de aplicar un cambio."""
    st.session_state.historial_versiones.append({
        "texto": texto,
        "bytes": bytes_doc,
        "ts": datetime.now().strftime("%H:%M:%S")
    })
    if len(st.session_state.historial_versiones) > 10:
        st.session_state.historial_versiones.pop(0)

def auditoria_tecnica_gastos(texto):
    """
    Realiza la auditoría de fechas para corregir errores de digitación
    y elimina filas sin registros temporales.
    """
    lineas = texto.split('\n')
    lineas_auditadas = []
    # Patrón para fechas (DD/MM/AAAA o similares)
    patron_fecha = r'(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})'
    
    for linea in lineas:
        fecha_match = re.search(patron_fecha, linea)
        if fecha_match:
            fecha_texto = fecha_match.group(1)
            # Corregir años truncados (ej: 202 -> 2024)
            partes = re.split(r'[/-]', fecha_texto)
            if len(partes) == 3 and len(partes[2]) == 3:
                partes[2] = "2024" 
                nueva_fecha = "/".join(partes)
                linea = linea.replace(fecha_texto, nueva_fecha)
            
            # Vinculación lógica: Solo aceptamos filas con contenido de valor (números/montos)
            if any(char.isdigit() for char in linea) and len(linea.split()) > 2:
                lineas_auditadas.append(linea)
                
    return "\n".join(lineas_auditadas)

# ══════════════════════════════════════════════════════════════
# FUNCIONES IA
# ══════════════════════════════════════════════════════════════
def solicitar_resumen_estructurado(texto):
    idioma_prompt = T("prompt_idioma")
    prompt = (
        f"Analista experto en documentos. {idioma_prompt} Devuelve SOLO JSON válido:\n"
        '{"titulo":"...","emoji_categoria":"📋","resumen_ejecutivo":"max 3 oraciones amigables",'
        '"metricas":{"Clave":"Valor"},"puntos_clave":["punto"],"hallazgo_destacado":"observación"}\n\n'
        f"DOCUMENTO:\n{texto[:12000]}"
    )
    r = llamar_ia(prompt)
    return extraer_json_seguro(r) if r else None

def extraer_cambio_con_regex(instruccion):
    patrones = [
        r"(?:cambia|reemplaza|sustituye|cambie)\s+['\"]?(.+?)['\"]?\s+(?:por|con|a)\s+['\"]?(.+?)['\"]?\s*$",
        r"['\"](.+?)['\"]\s*(?:→|->|=>|por|con)\s*['\"]?(.+?)['\"]?\s*$",
        r"(.+?)\s*(?:→|->|=>)\s*(.+)",
    ]
    for pat in patrones:
        m = re.search(pat,instruccion.strip(),re.IGNORECASE)
        if m:
            b=m.group(1).strip().strip("'\"")
            r2=m.group(2).strip().strip("'\"")
            if b and r2: return [{"buscar":b,"reemplazar":r2}]
    return []

def solicitar_cambios(instruccion, texto_doc=""):
    ctx = f"\n\nCONTENIDO DEL DOCUMENTO:\n{texto_doc[:4000]}" if texto_doc else ""
    idioma_prompt = T("prompt_idioma")
    prompt = (
        f"Asistente de edición de documentos. {idioma_prompt}\nINSTRUCCIÓN: \"{instruccion}\"{ctx}\n\n"
        "REGLAS:\n1. cambia X por Y → buscar=X, reemplazar=Y\n"
        "2. agrega DATO a PERSONA → buscar=PERSONA exacto, reemplazar='PERSONA DATO'\n"
        "3. completa campo de X con Y → buscar=X, reemplazar='X Y'\n"
        "4. SIEMPRE usa texto EXACTO del doc como buscar\n"
        "5. Si pide formato (negrita/mayúsculas/cursiva) → agrega campo 'formato'\n"
        "Responde SOLO JSON array:\n"
        '[{"buscar":"texto_exacto","reemplazar":"texto_nuevo"}]'
    )
    r = llamar_ia(prompt)
    if r:
        res = extraer_json_seguro(r,es_lista=True)
        if res and isinstance(res,list):
            v = [c for c in res if isinstance(c,dict) and "buscar" in c and "reemplazar" in c
                 and str(c["buscar"]).strip() and str(c["reemplazar"]).strip() and c["buscar"]!=c["reemplazar"]]
            if v: return v
    return extraer_cambio_con_regex(instruccion)

def preguntar_al_documento(pregunta, texto):
    ctx = "\n".join([f"{m['rol']}: {m['texto']}" for m in st.session_state.historial_chat[-6:]])
    idioma_prompt = T("prompt_idioma")
    prompt = (
        f"Asistente experto en documentos. {idioma_prompt}\nDOCUMENTO:\n{texto[:10000]}\n\n"
        f"CONVERSACIÓN:\n{ctx}\n\nPREGUNTA: {pregunta}\nResponde conciso y directo en español."
    )
    return llamar_ia(prompt) or "No pude procesar tu pregunta."

def detectar_anomalias(texto):
    idioma_prompt = T("prompt_idioma")
    prompt = (
        f"Analiza el documento. {idioma_prompt} Devuelve SOLO JSON:\n"
        '{"nivel_general":"Excelente/Bueno/Regular/Deficiente","puntaje":85,'
        '"criticos":["..."],"altos":["..."],"medios":["..."],"leves":["..."],'
        '"recomendacion":"..."}\n\n'
        f"DOCUMENTO:\n{texto[:12000]}"
    )
    r = llamar_ia(prompt)
    return extraer_json_seguro(r) if r else None

def detectar_tipo_imagen(texto_raw):
    lineas = [l for l in texto_raw.split('\n') if l.strip()]
    if not lineas: return "word"
    lineas_con_cols = sum(1 for l in lineas if len(re.split(r'\s{2,}|\t', l.strip())) >= 2)
    ratio_tabla = lineas_con_cols / max(len(lineas), 1)
    if ratio_tabla >= 0.5:
        return "excel"
    return "word"

def interpretar_imagen_documento(imagen_bytes, mime_type="image/jpeg", formato_salida="auto"):
    texto_raw = None
    try:
        from PIL import Image, ImageEnhance, ImageFilter
        import pytesseract, io
        img = Image.open(io.BytesIO(imagen_bytes))
        img_gray = img.convert('L')
        img_contrast = ImageEnhance.Contrast(img_gray).enhance(2.0)
        img_sharp = img_contrast.filter(ImageFilter.SHARPEN)
        texto_raw = pytesseract.image_to_string(img_sharp, lang='spa+eng', config='--psm 6')
        if not texto_raw.strip():
            texto_raw = None
    except:
        texto_raw = None

    if not texto_raw:
        try:
            import base64
            img_b64 = base64.b64encode(imagen_bytes).decode("utf-8")
            model = genai.GenerativeModel("gemini-1.5-flash")
            resp = model.generate_content([
                "Extrae TODO el texto de esta imagen. Si es una tabla, usa tabulaciones para separar columnas. No agregues comentarios.",
                {"mime_type": mime_type, "data": img_b64}
            ])
            texto_raw = resp.text
        except:
            return None, "word"

    tipo = detectar_tipo_imagen(texto_raw) if formato_salida == "auto" else formato_salida
    return texto_raw, tipo

# ══════════════════════════════════════════════════════════════
# EXPORTACIÓN
# ══════════════════════════════════════════════════════════════
def exportar_a_word(texto, titulo="Informe Corregido"):
    doc = Document()
    s = doc.styles['Normal']; s.font.name = 'Arial'; s.font.size = Pt(11)
    h = doc.add_heading(titulo, 0); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generado por Oro Asistente - {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph("_" * 40)
    for p in texto.split('\n'):
        if p.strip(): doc.add_paragraph(p.strip())
    out = BytesIO(); doc.save(out); return out.getvalue()

def exportar_a_excel(texto):
    wb = openpyxl.Workbook(); ws = wb.active; ws.title = "Datos"
    filas = [re.split(r'\s{2,}|\t|;', l.strip()) for l in texto.split('\n') if l.strip()]
    for r_idx, fila in enumerate(filas, 1):
        for c_idx, val in enumerate(fila, 1):
            cell = ws.cell(row=r_idx, column=c_idx, value=val)
            if r_idx == 1:
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill("solid", fgColor="4f46e5")
    out = BytesIO(); wb.save(out); return out.getvalue()

def exportar_a_pdf(texto):
    pdf = FPDF(); pdf.add_page(); pdf.set_font("Arial", size=10)
    for l in texto.split('\n'):
        pdf.multi_cell(0, 8, l.strip().encode('latin-1', 'replace').decode('latin-1'))
    return pdf.output(dest='S').encode('latin-1')

# ══════════════════════════════════════════════════════════════
# LÓGICA DE CARGA
# ══════════════════════════════════════════════════════════════
if st.session_state.modo_entrada == "archivo":
    fu = st.file_uploader(T("subir"), type=["docx", "xlsx", "pdf", "jpg", "png", "jpeg"], label_visibility="collapsed")
    if fu and fu.name != st.session_state.nombre_archivo:
        with st.spinner(T("cargando")):
            st.session_state.nombre_archivo = fu.name
            st.session_state.archivo_bytes = fu.getvalue()
            st.session_state.lista_cambios = []
            st.session_state.resumen_data = None
            st.session_state.guia_paso = 1
            ext = fu.name.split('.')[-1].lower()
            txt = ""
            if ext == "docx":
                st.session_state.archivo_tipo = "word"
                doc = Document(BytesIO(st.session_state.archivo_bytes))
                txt = "\n".join([p.text for p in doc.paragraphs])
            elif ext == "xlsx":
                st.session_state.archivo_tipo = "excel"
                wb = openpyxl.load_workbook(BytesIO(st.session_state.archivo_bytes), data_only=True)
                ws = wb.active
                txt = "\n".join(["\t".join([str(c.value or "") for c in r]) for r in ws.rows])
            elif ext == "pdf":
                st.session_state.archivo_tipo = "pdf"
                try:
                    pdf = PyPDF2.PdfReader(BytesIO(st.session_state.archivo_bytes))
                    txt = "\n".join([p.extract_text() for p in pdf.pages])
                except:
                    if PYMUPDF_OK:
                        doc = fitz.open(stream=st.session_state.archivo_bytes, filetype="pdf")
                        txt = "\n".join([p.get_text() for p in doc])
            elif ext in ["jpg","jpeg","png"]:
                txt, tipo = interpretar_imagen_documento(st.session_state.archivo_bytes, fu.type)
                st.session_state.archivo_tipo = tipo
            st.session_state.texto_extraido = txt
            st.session_state.texto_corregido = txt
            st.rerun()

# ══════════════════════════════════════════════════════════════
# DASHBOARD
# ══════════════════════════════════════════════════════════════
if st.session_state.texto_extraido:
    texto_activo = st.session_state.texto_corregido
    
    # Header del archivo
    ext_ico = {"word":"📄","excel":"📊","pdf":"📕"}.get(st.session_state.archivo_tipo,"📄")
    st.markdown(f"""<div class="file-badge">
        <div class="file-icon">{ext_ico}</div>
        <div style="flex:1">
            <div class="file-info-name">{st.session_state.nombre_archivo}</div>
            <div class="file-info-stats">
                <span class="stat-chip">{len(texto_activo.split())} {T('palabras')}</span>
                <span class="stat-chip">{len(st.session_state.lista_cambios)} {T('cambios')}</span>
                <span class="stat-chip">{len(st.session_state.historial_versiones)} {T('versiones')}</span>
            </div>
        </div>
    </div>""", unsafe_allow_html=True)

    # Botones principales
    c1, c2, c3 = st.columns(3)
    with c1:
        if st.button(T("analizar"), use_container_width=True, type="primary"):
            with st.spinner(T("analizando")):
                st.session_state.resumen_data = solicitar_resumen_estructurado(texto_activo)
                if st.session_state.resumen_data: st.session_state.guia_paso = 2
            st.rerun()
    with c2:
        if st.button(T("evaluar"), use_container_width=True):
            st.session_state.ejecutar_evaluacion = True
    with c3:
        with st.expander("📥"):
            if st.session_state.archivo_tipo == "excel":
                st.download_button("📊 Excel", exportar_a_excel(texto_activo), "auditado.xlsx", use_container_width=True)
            else:
                st.download_button("📄 Word", exportar_a_word(texto_activo), "auditado.docx", use_container_width=True)
            st.download_button("📕 PDF", exportar_a_pdf(texto_activo), "auditado.pdf", use_container_width=True)

    # Evaluación de Calidad
    if st.session_state.get("ejecutar_evaluacion"):
        with st.spinner(T("evaluando")):
            eval_data = detectar_anomalias(texto_activo)
            if eval_data:
                st.markdown(f"### Calidad del Documento: {eval_data['nivel_general']}")
                st.progress(eval_data['puntaje']/100)
                if eval_data['criticos']: st.error("Críticos: " + ", ".join(eval_data['criticos']))
                if eval_data['altos']: st.warning("Altos: " + ", ".join(eval_data['altos']))
                st.info("Recomendación: " + eval_data['recomendacion'])
        if st.button(T("cerrar_eval")):
            st.session_state.ejecutar_evaluacion = False; st.rerun()

    # Resumen Inteligente
    if st.session_state.resumen_data:
        rd = st.session_state.resumen_data
        st.markdown(f"""<div class="summary-card">
            <div class="summary-card-title">{rd.get('emoji_categoria','📋')} {rd.get('titulo','Resumen')}</div>
            {rd.get('resumen_ejecutivo','')}
        </div>""", unsafe_allow_html=True)
        
        m_cols = st.columns(len(rd.get('metricas',{})))
        for i, (k,v) in enumerate(rd.get('metricas',{}).items()):
            with m_cols[i]:
                st.markdown(f"""<div class="metric-pill">
                    <div class="metric-pill-label">{k}</div>
                    <div class="metric-pill-value">{v}</div>
                </div>""", unsafe_allow_html=True)
        
        st.markdown(f"""<div class="hallazgo-card">
            <span style="font-weight:800;color:#6b83f8">{T('hallazgo')}</span> {rd.get('hallazgo_destacado','')}
        </div>""", unsafe_allow_html=True)

    # CHAT / EDICIÓN
    st.markdown(f"<div class='section-title'>💬 {T('chat_titulo')}</div>", unsafe_allow_html=True)
    
    for msg in st.session_state.historial_chat:
        with st.chat_message("user" if msg["rol"]=="Usuario" else "assistant"):
            st.markdown(msg["texto"])

    entrada = st.chat_input(T("chat_placeholder"))
    if entrada:
        st.session_state.historial_chat.append({"rol":"Usuario","texto":entrada})
        
        # LÓGICA DE AUDITORÍA INCORPORADA
        if any(x in entrada.lower() for x in ["auditoria", "auditoría", "corregir fechas", "limpiar registros"]):
            with st.spinner("Ejecutando auditoría técnica..."):
                guardar_version(st.session_state.texto_corregido, st.session_state.archivo_bytes)
                resultado_audit = auditoria_tecnica_gastos(texto_activo)
                st.session_state.texto_corregido = resultado_audit
                msg = "✅ Auditoría completada. He corregido errores de digitación en fechas y eliminado las filas sin registros temporales según lo recomendado."
                st.session_state.historial_chat.append({"rol":"Asistente","texto":msg})
        
        # Edición por lenguaje natural
        elif any(x in entrada.lower() for x in ["cambia","reemplaza","sustituye","→","->","agrega"]):
            with st.spinner(T("pensando")):
                cambios = solicitar_cambios(entrada, texto_activo)
            if cambios:
                guardar_version(st.session_state.texto_corregido, st.session_state.archivo_bytes)
                nuevo_texto = texto_activo
                realizados = []
                for c in cambios:
                    busc, reem = str(c["buscar"]), str(c["reemplazar"])
                    if busc in nuevo_texto:
                        nuevo_texto = nuevo_texto.replace(busc, reem)
                        realizados.append(f"'{busc}' → '{reem}'")
                st.session_state.texto_corregido = nuevo_texto
                st.session_state.lista_cambios.extend(realizados)
                msg = f"✅ Aplicado: {', '.join(realizados)}. Revisa la vista previa arriba y confirma."
                st.session_state.historial_chat.append({"rol":"Asistente","texto":msg})
            else:
                msg = "No encontré qué cambiar. Intenta: *cambia 'palabra original' por 'palabra nueva'*"
                st.session_state.historial_chat.append({"rol":"Asistente","texto":msg})
        else:
            with st.spinner(T("pensando")):
                resp = preguntar_al_documento(entrada, texto_activo)
            st.session_state.historial_chat.append({"rol":"Asistente","texto":resp})
        st.rerun()

else:
    st.markdown("""<div class="empty-state">
        <div class="empty-icon">📂</div>
        <div class="empty-title">Sube un archivo para empezar</div>
        <div class="empty-hint">
            Analiza documentos con IA · Edita con lenguaje natural<br>
            Exporta a Word, Excel o PDF
        </div>
        <div class="format-badges">
            <span class="format-badge">.docx</span>
            <span class="format-badge">.xlsx</span>
            <span class="format-badge">.pdf</span>
            <span class="format-badge">📷 foto</span>
        </div>
    </div>""", unsafe_allow_html=True)

zona_horaria = pytz.timezone('America/Caracas')
hora = datetime.now(zona_horaria).strftime("%H:%M")
st.markdown(f"<div class='oro-footer'>Oro Asistente v3.1 · {hora} · Desarrollado para Auditoría Deportiva</div>", unsafe_allow_html=True)
